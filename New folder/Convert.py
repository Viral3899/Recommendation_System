import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from docx import Document
import pytesseract
from PIL import Image
from pdf2image import convert_from_path

# Download stopwords (if not already downloaded)
nltk.download('stopwords')

# Function to convert PDF to images
def convert_pdf_to_images(pdf_path):
    return convert_from_path(pdf_path, poppler_path=r'C:\poppler\Library\bin')

# PDF file path
pdf_path = r'C:\Users\viral\Downloads\documents\documents\123.pdf'

# Convert PDF to images
images = convert_pdf_to_images(pdf_path)

# Extract text from each image using OCR (Tesseract)
extracted_text = []
for image in images:
    extracted_text.append(pytesseract.image_to_string(image))


# Combine extracted text from all images
bill_text = ' '.join(extracted_text)

# Tokenize the text into words
words = word_tokenize(bill_text)

# Remove stopwords (common words that do not carry much meaning)
stop_words = set(stopwords.words('english'))
filtered_words = [word for word in words if word.casefold() not in stop_words]

# Extracting patient name
patient_name = None
for i in range(len(filtered_words)):
    if filtered_words[i].casefold() == 'patient' and filtered_words[i+1].casefold() == 'name':
        patient_name = filtered_words[i+2]
        break

# Extracting date
date = None
for i in range(len(filtered_words)):
    if filtered_words[i].casefold() == 'date':
        date = filtered_words[i+1]
        break

# Extracting total amount
total_amount = None
for i in range(len(filtered_words)):
    if filtered_words[i].casefold() == 'total' and filtered_words[i+1].casefold() == 'amount':
        total_amount = filtered_words[i+2]
        break

# Create a new Word document
document = Document()

# Add content to the document
document.add_heading('Hospital Report Information', level=1)
document.add_paragraph(f'Patient Name: {patient_name}')
document.add_paragraph(f'Date: {date}')
document.add_paragraph(f'Total Amount: {total_amount}')

# Save the document
print('saving........')
document.save('hospital_report.docx')
