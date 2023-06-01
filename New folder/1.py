import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Function to convert PDF to images
def convert_pdf_to_images(pdf_path):
    return convert_from_path(pdf_path)

# Function to extract text from image using OCR (Tesseract)
def extract_text_from_image(image):
    return pytesseract.image_to_string(image)

# PDF file path
pdf_path = r'C:\Users\viral\Downloads\documents\documents\File-2 Admitted and Discharge Summary.pdf'

# Convert PDF to images
images = convert_pdf_to_images(pdf_path)

# Create a new Word document
document = Document()

# Process each image and extract text
for image in images:
    # Extract text from the image
    extracted_text = extract_text_from_image(image)

    # Save the image to a BytesIO object
    img_io = BytesIO()
    image.save(img_io, format='PNG')
    img_io.seek(0)

    # Create a new paragraph in the Word document
    paragraph = document.add_paragraph()

    # Add the image to the paragraph
    run = paragraph.add_run()
    run.add_picture(img_io, width=Inches(6))

    # Add the extracted text to the paragraph
    paragraph.add_run("\n")
    paragraph.add_run(extracted_text)

# Save the document
document.save('extracted_text2.docx')

print('PDF text extracted and saved to extracted_text.docx successfully.')


